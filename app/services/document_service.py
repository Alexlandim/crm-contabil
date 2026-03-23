from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from sqlalchemy.orm import Session

from app.models import CompanySetting, Proposal

PRIMARY_COLOR = colors.HexColor("#193A6A")
SECONDARY_COLOR = colors.HexColor("#EAF2FF")
ACCENT_COLOR = colors.HexColor("#244B84")
MUTED_TEXT = colors.HexColor("#4B5563")
LIGHT_BORDER = colors.HexColor("#D7DFEA")
TOTAL_FILL = colors.HexColor("#F1F6FF")

DEFAULT_SECTION_TEXTS = {
    "presentation": (
        "A Contécnica Contabilidade apresenta esta proposta com foco em atendimento consultivo, "
        "confiabilidade operacional e suporte técnico contínuo, buscando entregar segurança nas rotinas "
        "contábeis, fiscais e trabalhistas, além de informações gerenciais que apoiem a tomada de decisão."
    ),
    "methodology": (
        "Nossa metodologia considera diagnóstico inicial, parametrização operacional, definição de fluxo de "
        "documentos, acompanhamento periódico e suporte recorrente, com atuação orientada por prazos, "
        "conformidade legal e clareza na comunicação com o cliente."
    ),
    "services_description": (
        "Os serviços serão executados conforme o escopo contratado, contemplando as rotinas compatíveis com a "
        "necessidade operacional da empresa, bem como a emissão de relatórios e orientações necessárias ao "
        "cumprimento das obrigações acessórias e principais."
    ),
    "extra_services": (
        "Serviços extraordinários, demandas fora do escopo recorrente, regularizações específicas, levantamentos "
        "retroativos, atendimentos presenciais não programados e projetos especiais poderão ser orçados em "
        "separado, mediante prévia validação do cliente."
    ),
    "general_conditions": (
        "Esta proposta foi elaborada com base nas informações fornecidas até a presente data. Eventuais alterações "
        "de escopo, volume operacional, regime tributário, quadro de colaboradores ou exigências legais supervenientes "
        "poderão ensejar reavaliação comercial."
    ),
    "closing": (
        "Permanecemos à disposição para quaisquer esclarecimentos adicionais e para a realização de reunião de "
        "apresentação, caso necessário. Será uma satisfação apoiar a estruturação e o crescimento do seu negócio."
    ),
}


def brl(value: Any) -> str:
    numeric_value = float(value or 0)
    return f"R$ {numeric_value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def normalize_multiline_text(text: str | None) -> list[str]:
    if not text:
        return []
    lines = [line.strip() for line in str(text).replace("\r", "").split("\n")]
    return [line for line in lines if line]


def get_default_template_texts() -> dict[str, str]:
    return DEFAULT_SECTION_TEXTS.copy()


def get_document_base_content(proposal: Proposal) -> dict[str, str]:
    snapshot = proposal.document_base
    return {
        "template_name": snapshot.template_name if snapshot else "Base padrão",
        "presentation": snapshot.presentation if snapshot and snapshot.presentation else DEFAULT_SECTION_TEXTS["presentation"],
        "methodology": snapshot.methodology if snapshot and snapshot.methodology else DEFAULT_SECTION_TEXTS["methodology"],
        "services_description": snapshot.services_description if snapshot and snapshot.services_description else DEFAULT_SECTION_TEXTS["services_description"],
        "extra_services": snapshot.extra_services if snapshot and snapshot.extra_services else DEFAULT_SECTION_TEXTS["extra_services"],
        "general_conditions": snapshot.general_conditions if snapshot and snapshot.general_conditions else DEFAULT_SECTION_TEXTS["general_conditions"],
        "closing": snapshot.closing if snapshot and snapshot.closing else DEFAULT_SECTION_TEXTS["closing"],
    }


def get_template_path() -> Path | None:
    primary = Path("templates_docx/template_proposta.docx")
    fallback = Path("templates_docx/template_proposta.docx.docx")

    if primary.exists():
        return primary
    if fallback.exists():
        return fallback
    return None


def section_has_graphics(part) -> bool:
    xml = part._element.xml
    return any(tag in xml for tag in ["w:drawing", "v:shape", "pic:pic", "wp:anchor", "wp:inline"])


def template_has_visual_header_or_footer(doc: Document) -> bool:
    try:
        for section in doc.sections:
            if section_has_graphics(section.header):
                return True
            if section_has_graphics(section.footer):
                return True
    except Exception:
        pass
    return False


def set_cell_text(cell, text, bold=False, align="left", font_size=10):
    cell.text = str(text or "")
    for paragraph in cell.paragraphs:
        if align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for run in paragraph.runs:
            run.font.name = "Arial"
            if run._element.rPr is None:
                run._element.get_or_add_rPr()
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            run.font.size = Pt(font_size)
            run.bold = bold

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clear_document_body(doc: Document):
    body = doc._element.body
    elements_to_remove = []

    for child in body.iterchildren():
        if child.tag != qn("w:sectPr"):
            elements_to_remove.append(child)

    for element in elements_to_remove:
        body.remove(element)


def set_document_margins(doc: Document, has_letterhead=False):
    section = doc.sections[0]

    if has_letterhead:
        section.top_margin = Cm(4.8)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(0.8)
    else:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)


def set_default_styles(doc: Document):
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal_style.font.size = Pt(10.5)

    if "Title" in doc.styles:
        title_style = doc.styles["Title"]
        title_style.font.name = "Arial"
        title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        title_style.font.size = Pt(20)
        title_style.font.bold = True

    if "Heading 1" in doc.styles:
        h1 = doc.styles["Heading 1"]
        h1.font.name = "Arial"
        h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        h1.font.size = Pt(13)
        h1.font.bold = True

    if "Heading 2" in doc.styles:
        h2 = doc.styles["Heading 2"]
        h2.font.name = "Arial"
        h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        h2.font.size = Pt(11.5)
        h2.font.bold = True


def add_header_footer(doc: Document, company: CompanySetting | None):
    section = doc.sections[0]

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if not header_paragraph.text.strip():
        header_run = header_paragraph.add_run(company.company_name if company else "Contécnica Contabilidade")
        header_run.font.name = "Arial"
        header_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        header_run.font.size = Pt(9)
        header_run.font.bold = True

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not footer_paragraph.text.strip():
        footer_run = footer_paragraph.add_run(
            company.proposal_footer if company and company.proposal_footer else "Documento gerado pelo CRM Pro."
        )
        footer_run.font.name = "Arial"
        footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        footer_run.font.size = Pt(8.5)


def add_spacer(doc: Document, height_pt: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    run = p.add_run("")
    run.font.size = Pt(height_pt)
    return p


def add_centered_paragraph(doc: Document, text: str, size=10.5, bold=False, color_rgb=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)

    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    return p


def add_section_heading(doc: Document, code: str, title: str):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{code}. {title.upper()}")
    run.bold = True
    return p


def add_multiline_block(doc: Document, text: str):
    lines = normalize_multiline_text(text)
    if not lines:
        lines = ["Não informado."]

    for line in lines:
        p = doc.add_paragraph()
        p.style = "Normal"
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.2
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(line)


def add_highlight_box(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    table.style = "Table Grid"

    for label, value in rows:
        row = table.add_row().cells
        set_cell_text(row[0], label, bold=True)
        set_cell_text(row[1], value)
        shade_cell(row[0], "EAF2FF")

    doc.add_paragraph("")


def add_client_information_table(doc: Document, proposal: Proposal):
    data = [
        ("Cliente", proposal.customer.legal_name if proposal.customer else ""),
        ("CNPJ", proposal.cnpj or "-"),
        ("Regime tributário", proposal.tax_regime or "-"),
        ("Segmento", proposal.business_segment or "-"),
        ("Quantidade de funcionários", str(proposal.employee_count or 0)),
        ("Média de notas fiscais", str(proposal.monthly_invoices_avg or 0)),
        ("Faturamento", proposal.faturamento or "-"),
        ("Média de faturamento mensal", brl(proposal.monthly_revenue_avg or 0)),
        ("Número da proposta", proposal.proposal_number or "-"),
        ("Emissão", proposal.issue_date.strftime("%d/%m/%Y") if proposal.issue_date else "-"),
        ("Validade", proposal.validity_date.strftime("%d/%m/%Y") if proposal.validity_date else "-"),
    ]
    add_highlight_box(doc, data)


def add_items_table_docx(doc: Document, proposal: Proposal):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = True

    headers = table.rows[0].cells
    set_cell_text(headers[0], "Descrição", bold=True, align="center")
    set_cell_text(headers[1], "Qtd.", bold=True, align="center")
    set_cell_text(headers[2], "Vlr. Unit.", bold=True, align="center")
    set_cell_text(headers[3], "Desconto", bold=True, align="center")
    set_cell_text(headers[4], "Total", bold=True, align="center")

    for cell in headers:
        shade_cell(cell, "D9EAF7")

    for item in proposal.items:
        row = table.add_row().cells
        set_cell_text(row[0], item.description or "")
        set_cell_text(row[1], item.quantity or 0, align="center")
        set_cell_text(row[2], brl(item.unit_price or 0), align="right")
        set_cell_text(row[3], brl(item.discount_amount or 0), align="right")
        set_cell_text(row[4], brl(item.line_total or 0), align="right")

    doc.add_paragraph("")

    summary_table = doc.add_table(rows=3, cols=2)
    summary_table.style = "Table Grid"
    summary_table.autofit = True

    summary_rows = [
        ("Subtotal", brl(proposal.subtotal_amount or 0)),
        ("Desconto global", brl(proposal.global_discount or 0)),
        ("Total da proposta", brl(proposal.total_amount or 0)),
    ]

    for index, (label, value) in enumerate(summary_rows):
        row = summary_table.rows[index].cells
        set_cell_text(row[0], label, bold=True)
        set_cell_text(row[1], value, bold=(index == 2), align="right")
        shade_cell(row[0], "F5F8FC" if index < 2 else "EAF2FF")
        if index == 2:
            shade_cell(row[1], "EAF2FF")

    doc.add_paragraph("")


def add_word_cover(doc: Document, company: CompanySetting | None, proposal: Proposal, has_letterhead=False):
    if has_letterhead:
        add_spacer(doc, 24)
        add_spacer(doc, 24)

    add_centered_paragraph(doc, "PROPOSTA COMERCIAL", size=21, bold=True, space_after=2)
    add_centered_paragraph(doc, "Prestação de Serviços Contábeis", size=14, bold=False, space_after=10)

    cover_box = doc.add_table(rows=4, cols=2)
    cover_box.style = "Table Grid"
    cover_box.autofit = True
    cover_data = [
        ("Empresa", proposal.customer.legal_name if proposal.customer else "Cliente"),
        ("Proposta", proposal.proposal_number or "-"),
        ("Emissão", proposal.issue_date.strftime("%d/%m/%Y") if proposal.issue_date else "-"),
        ("Validade", proposal.validity_date.strftime("%d/%m/%Y") if proposal.validity_date else "-"),
    ]

    for idx, (label, value) in enumerate(cover_data):
        row = cover_box.rows[idx].cells
        set_cell_text(row[0], label, bold=True)
        set_cell_text(row[1], value)
        shade_cell(row[0], "EAF2FF")

    doc.add_paragraph("")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    intro.paragraph_format.line_spacing = 1.2
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "Apresentamos a seguir nossa proposta de prestação de serviços, elaborada com base nas informações "
        "compartilhadas até o momento e estruturada para proporcionar clareza comercial, segurança operacional "
        "e visão executiva do escopo proposto."
    )

    doc.add_page_break()


def add_word_index(doc: Document):
    add_centered_paragraph(doc, "SUMÁRIO EXECUTIVO", size=15, bold=True, space_after=10)

    items = [
        "01 - Apresentação da Contécnica",
        "02 - Metodologia de trabalho",
        "03 - Descrição dos serviços prestados",
        "04 - Informações do cliente",
        "05 - Investimento e composição comercial",
        "06 - Serviços extraordinários",
        "07 - Condições gerais",
        "08 - Encerramento",
    ]

    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(2)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(item)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(11)
        run.bold = True

    doc.add_page_break()


def export_proposal_docx(db: Session, proposal: Proposal, output_path: str):
    company = db.query(CompanySetting).first()
    base_content = get_document_base_content(proposal)
    template_path = get_template_path()

    if template_path:
        doc = Document(template_path)
        has_letterhead = template_has_visual_header_or_footer(doc)
        clear_document_body(doc)
    else:
        doc = Document()
        has_letterhead = False

    set_document_margins(doc, has_letterhead=has_letterhead)
    set_default_styles(doc)

    if not has_letterhead:
        add_header_footer(doc, company)

    add_word_cover(doc, company, proposal, has_letterhead=has_letterhead)
    add_word_index(doc)

    add_section_heading(doc, "01", "Apresentação da Contécnica")
    add_multiline_block(doc, base_content["presentation"])

    add_section_heading(doc, "02", "Metodologia de Trabalho")
    add_multiline_block(doc, base_content["methodology"])

    add_section_heading(doc, "03", "Descrição dos Serviços Prestados")
    add_multiline_block(doc, base_content["services_description"])

    add_section_heading(doc, "04", "Informações do Cliente")
    add_client_information_table(doc, proposal)

    add_section_heading(doc, "05", "Investimento e Composição Comercial")
    add_items_table_docx(doc, proposal)

    if proposal.notes:
        p = doc.add_paragraph()
        p.style = "Heading 2"
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        p.add_run("Observações complementares")
        add_multiline_block(doc, proposal.notes)

    add_section_heading(doc, "06", "Serviços Extraordinários")
    add_multiline_block(doc, base_content["extra_services"])

    add_section_heading(doc, "07", "Condições Gerais")
    add_multiline_block(doc, base_content["general_conditions"])

    add_section_heading(doc, "08", "Encerramento")
    add_multiline_block(doc, base_content["closing"])

    doc.save(output_path)


def get_pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "PdfTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=24,
            alignment=TA_CENTER,
            textColor=PRIMARY_COLOR,
            spaceAfter=6,
        ),
        "subtitle": ParagraphStyle(
            "PdfSubtitle",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=11.5,
            leading=15,
            alignment=TA_CENTER,
            textColor=MUTED_TEXT,
            spaceAfter=16,
        ),
        "section": ParagraphStyle(
            "PdfSection",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11.5,
            leading=14,
            textColor=PRIMARY_COLOR,
            spaceBefore=10,
            spaceAfter=8,
        ),
        "body": ParagraphStyle(
            "PdfBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            alignment=TA_JUSTIFY,
            textColor=colors.black,
            spaceAfter=6,
        ),
        "small": ParagraphStyle(
            "PdfSmall",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=10,
            textColor=MUTED_TEXT,
            alignment=TA_CENTER,
        ),
        "meta_label": ParagraphStyle(
            "PdfMetaLabel",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8.8,
            leading=10,
            textColor=PRIMARY_COLOR,
            alignment=TA_LEFT,
        ),
        "meta_value": ParagraphStyle(
            "PdfMetaValue",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=12,
            textColor=colors.black,
            alignment=TA_LEFT,
        ),
        "table": ParagraphStyle(
            "PdfTable",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.2,
            leading=11,
            textColor=colors.black,
            alignment=TA_LEFT,
        ),
        "table_right": ParagraphStyle(
            "PdfTableRight",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.2,
            leading=11,
            textColor=colors.black,
            alignment=TA_RIGHT,
        ),
    }


def ptext(text: str, style: ParagraphStyle) -> Paragraph:
    safe = (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(safe.replace("\n", "<br/>"), style)


def build_pdf_meta_table(proposal: Proposal, styles: dict[str, ParagraphStyle]) -> Table:
    data = [
        [ptext("Cliente", styles["meta_label"]), ptext(proposal.customer.legal_name if proposal.customer else "-", styles["meta_value"]), ptext("Proposta", styles["meta_label"]), ptext(proposal.proposal_number or "-", styles["meta_value"])],
        [ptext("CNPJ", styles["meta_label"]), ptext(proposal.cnpj or "-", styles["meta_value"]), ptext("Emissão", styles["meta_label"]), ptext(proposal.issue_date.strftime("%d/%m/%Y") if proposal.issue_date else "-", styles["meta_value"])],
        [ptext("Regime", styles["meta_label"]), ptext(proposal.tax_regime or "-", styles["meta_value"]), ptext("Validade", styles["meta_label"]), ptext(proposal.validity_date.strftime("%d/%m/%Y") if proposal.validity_date else "-", styles["meta_value"])],
        [ptext("Segmento", styles["meta_label"]), ptext(proposal.business_segment or "-", styles["meta_value"]), ptext("Total", styles["meta_label"]), ptext(brl(proposal.total_amount or 0), styles["meta_value"])],
    ]

    table = Table(data, colWidths=[2.2 * cm, 6.0 * cm, 2.2 * cm, 5.4 * cm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), SECONDARY_COLOR),
                ("BACKGROUND", (2, 0), (2, -1), SECONDARY_COLOR),
                ("BOX", (0, 0), (-1, -1), 0.6, LIGHT_BORDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, LIGHT_BORDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    return table


def build_pdf_client_table(proposal: Proposal, styles: dict[str, ParagraphStyle]) -> Table:
    rows = [
        ("Cliente", proposal.customer.legal_name if proposal.customer else "-"),
        ("CNPJ", proposal.cnpj or "-"),
        ("Regime tributário", proposal.tax_regime or "-"),
        ("Segmento", proposal.business_segment or "-"),
        ("Qtd. de funcionários", str(proposal.employee_count or 0)),
        ("Média de notas fiscais", str(proposal.monthly_invoices_avg or 0)),
        ("Faturamento", proposal.faturamento or "-"),
        ("Média de faturamento mensal", brl(proposal.monthly_revenue_avg or 0)),
    ]
    data = [[ptext(label, styles["meta_label"]), ptext(value, styles["meta_value"])] for label, value in rows]
    table = Table(data, colWidths=[5.4 * cm, 11.0 * cm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), SECONDARY_COLOR),
                ("BOX", (0, 0), (-1, -1), 0.6, LIGHT_BORDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, LIGHT_BORDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return table


def build_pdf_items_table(proposal: Proposal, styles: dict[str, ParagraphStyle]) -> Table:
    data = [[
        ptext("Descrição", styles["meta_label"]),
        ptext("Qtd.", styles["meta_label"]),
        ptext("Vlr. Unit.", styles["meta_label"]),
        ptext("Desconto", styles["meta_label"]),
        ptext("Total", styles["meta_label"]),
    ]]

    for item in proposal.items:
        data.append([
            ptext(item.description or "", styles["table"]),
            ptext(str(item.quantity or 0), styles["table_right"]),
            ptext(brl(item.unit_price or 0), styles["table_right"]),
            ptext(brl(item.discount_amount or 0), styles["table_right"]),
            ptext(brl(item.line_total or 0), styles["table_right"]),
        ])

    table = Table(data, colWidths=[7.8 * cm, 1.5 * cm, 2.4 * cm, 2.5 * cm, 2.6 * cm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), SECONDARY_COLOR),
                ("TEXTCOLOR", (0, 0), (-1, 0), PRIMARY_COLOR),
                ("BOX", (0, 0), (-1, -1), 0.6, LIGHT_BORDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, LIGHT_BORDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return table


def build_pdf_summary_table(proposal: Proposal, styles: dict[str, ParagraphStyle]) -> Table:
    data = [
        [ptext("Subtotal", styles["meta_label"]), ptext(brl(proposal.subtotal_amount or 0), styles["table_right"])],
        [ptext("Desconto global", styles["meta_label"]), ptext(brl(proposal.global_discount or 0), styles["table_right"])],
        [ptext("Total da proposta", styles["meta_label"]), ptext(brl(proposal.total_amount or 0), styles["table_right"])],
    ]
    table = Table(data, colWidths=[5.5 * cm, 4.5 * cm], hAlign="RIGHT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), SECONDARY_COLOR),
                ("BACKGROUND", (0, 2), (-1, 2), TOTAL_FILL),
                ("BOX", (0, 0), (-1, -1), 0.6, LIGHT_BORDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, LIGHT_BORDER),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return table


def pdf_header_footer(canvas, doc, company: CompanySetting | None):
    canvas.saveState()
    width, height = A4

    canvas.setStrokeColor(PRIMARY_COLOR)
    canvas.setLineWidth(1)
    canvas.line(doc.leftMargin, height - 1.8 * cm, width - doc.rightMargin, height - 1.8 * cm)

    canvas.setFont("Helvetica-Bold", 11)
    canvas.setFillColor(PRIMARY_COLOR)
    canvas.drawString(doc.leftMargin, height - 1.35 * cm, company.company_name if company else "Contécnica Contabilidade")

    contact_parts = []
    if company and company.company_phone:
        contact_parts.append(company.company_phone)
    if company and company.company_email:
        contact_parts.append(company.company_email)
    if company and company.company_site:
        contact_parts.append(company.company_site)
    header_right = " | ".join(contact_parts) if contact_parts else "Proposta comercial"
    canvas.setFont("Helvetica", 8.5)
    canvas.setFillColor(MUTED_TEXT)
    canvas.drawRightString(width - doc.rightMargin, height - 1.32 * cm, header_right)

    canvas.setStrokeColor(LIGHT_BORDER)
    canvas.line(doc.leftMargin, 1.45 * cm, width - doc.rightMargin, 1.45 * cm)
    footer_text = company.proposal_footer if company and company.proposal_footer else "Documento gerado pelo CRM Pro."
    canvas.setFont("Helvetica", 8)
    canvas.drawString(doc.leftMargin, 1.0 * cm, footer_text[:110])
    canvas.drawRightString(width - doc.rightMargin, 1.0 * cm, f"Página {canvas.getPageNumber()}")
    canvas.restoreState()


def export_proposal_pdf(db: Session, proposal: Proposal, output_path: str):
    company = db.query(CompanySetting).first()
    base_content = get_document_base_content(proposal)
    styles = get_pdf_styles()

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=2.0 * cm,
        leftMargin=2.0 * cm,
        topMargin=2.8 * cm,
        bottomMargin=2.2 * cm,
        title=f"Proposta {proposal.proposal_number}",
        author=company.company_name if company else "Contécnica Contabilidade",
    )

    story = []
    story.append(ptext("PROPOSTA COMERCIAL", styles["title"]))
    story.append(ptext("Prestação de Serviços Contábeis", styles["subtitle"]))
    story.append(build_pdf_meta_table(proposal, styles))
    story.append(Spacer(1, 0.45 * cm))
    story.append(HRFlowable(width="100%", thickness=0.8, color=LIGHT_BORDER))
    story.append(Spacer(1, 0.35 * cm))
    story.append(
        ptext(
            "Apresentamos a seguir uma proposta estruturada para atendimento contábil, fiscal e consultivo, "
            "organizada de forma objetiva para facilitar a avaliação comercial e técnica do escopo ofertado.",
            styles["body"],
        )
    )
    story.append(PageBreak())

    story.append(ptext("01. APRESENTAÇÃO DA CONTÉCNICA", styles["section"]))
    story.append(ptext(base_content["presentation"], styles["body"]))

    story.append(ptext("02. METODOLOGIA DE TRABALHO", styles["section"]))
    story.append(ptext(base_content["methodology"], styles["body"]))

    story.append(ptext("03. DESCRIÇÃO DOS SERVIÇOS PRESTADOS", styles["section"]))
    story.append(ptext(base_content["services_description"], styles["body"]))

    story.append(ptext("04. INFORMAÇÕES DO CLIENTE", styles["section"]))
    story.append(build_pdf_client_table(proposal, styles))
    story.append(Spacer(1, 0.25 * cm))

    story.append(ptext("05. INVESTIMENTO E COMPOSIÇÃO COMERCIAL", styles["section"]))
    story.append(build_pdf_items_table(proposal, styles))
    story.append(Spacer(1, 0.25 * cm))
    story.append(build_pdf_summary_table(proposal, styles))

    if proposal.notes:
        story.append(Spacer(1, 0.3 * cm))
        story.append(ptext("Observações complementares", styles["section"]))
        story.append(ptext(proposal.notes, styles["body"]))

    story.append(ptext("06. SERVIÇOS EXTRAORDINÁRIOS", styles["section"]))
    story.append(ptext(base_content["extra_services"], styles["body"]))

    story.append(ptext("07. CONDIÇÕES GERAIS", styles["section"]))
    story.append(ptext(base_content["general_conditions"], styles["body"]))

    story.append(ptext("08. ENCERRAMENTO", styles["section"]))
    story.append(
        KeepTogether(
            [
                ptext(base_content["closing"], styles["body"]),
                Spacer(1, 0.4 * cm),
                ptext(company.company_name if company else "Contécnica Contabilidade", styles["meta_label"]),
            ]
        )
    )

    doc.build(story, onFirstPage=lambda c, d: pdf_header_footer(c, d, company), onLaterPages=lambda c, d: pdf_header_footer(c, d, company))
